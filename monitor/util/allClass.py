# coding:utf-8
# 所有功能点类
import time, datetime, shutil, os, webbrowser, sys, re, json
import sqlite3 as sl
import pyautogui, win32gui, win32con, win32clipboard
import docx, xlrd
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
import paramiko
from PIL import Image, ImageDraw, ImageFont
from urllib.request import Request
from urllib.request import urlopen
from lxml import etree
from monitor.otherModels import *
from monitor.util.preload import logger
from jpype import *


class WordFunctions:

    def __init__(self):
        pass

    # 移除word段落中的文本
    def word_dl_replace_text(self, word_path, old_text, new_text):
        doc = docx.Document(word_path)
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in inline:
                    if old_text in i.text:
                        text = i.text.replace(old_text, new_text)
                        i.text = text
        try:
            doc.save(word_path)
            return True
        except PermissionError:
            logger.exception('文件无法保存，请确认是否打开状态！')
            return False
        except Exception:
            logger.exception('未知错误！')
            return False

    # 移除word段落中的图片(标记文本会被替换掉)
    def word_dl_replace_pic(self, word_path, tag_text, pic):
        doc = docx.Document(word_path)
        for p in doc.paragraphs:
            if tag_text in p.text:
                inline = p.runs
                for i in inline:
                    if tag_text in i.text:
                        text = i.text.replace(tag_text, '')
                        i.text = text
                        i.add_picture(pic)
        try:
            doc.save(word_path)
            return True
        except PermissionError:
            logger.exception('文件无法保存，请确认是否打开状态！')
            return False
        except Exception:
            logger.exception('未知错误！')
            return False

    # 移除word表格中的文本,默认选择第一个表格
    def word_table_replace_text(self, word_path, old_text, new_text, chose=0):
        chose = int(chose)
        doc = docx.Document(word_path)
        for row in doc.tables[chose].rows:
            for cell in row.cells:
                if old_text in cell.text:
                    text = cell.text.replace(old_text, new_text)
                    cell.text = text
        try:
            doc.save(word_path)
            return True
        except PermissionError:
            logger.exception('文件无法保存，请确认是否打开状态！')
            return False
        except Exception:
            logger.exception('未知错误！')
            return False

    # 移除word段落中的图片(标记文本会被替换掉) 只支持PNG图片
    def word_table_replace_pic(self, word_path, tag_text, png_pic, pic_size='1', chose='0'):
        pic_size = int(pic_size)
        chose = int(chose)
        doc = docx.Document(word_path)
        for row in doc.tables[chose].rows:
            for cell in row.cells:
                if tag_text in cell.text:
                    cell.text = cell.text.replace(tag_text, '')
                    run = cell.paragraphs[0].runs
                    run[0].add_picture(png_pic, width=Inches(pic_size))
        try:
            doc.save(word_path)
            return True
        except PermissionError:
            logger.exception('文件无法保存，请确认是否打开状态！')
            return False
        except Exception:
            logger.exception('未知错误！')
            return False

    # 批量更新word段落中的内容,默认选择第一个表格,row_start开始行，col_start开始列
    def word_table_new_content(self, word_path, data, row_start=0, col_start=0, chose=0):
        # data格式 [['1', 0, '人员信息数据处理'], ['2', 0, '人员信息数据处理']]
        row_start = int(row_start)
        col_start = int(col_start)
        chose = int(chose)
        doc = docx.Document(word_path)
        for rowid, row in enumerate(doc.tables[chose].rows[row_start:]):
            for cellid, cell in enumerate(row.cells[col_start:]):
                cell.text = ''  # 先全部清空
                if data[rowid][cellid] == None:
                    # cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.DECIMAL
                    run = cell.paragraphs[0].add_run(None)
                    run.font.name = '宋体'
                    run.font.size = 120000
                else:
                    # cell.paragraphs[0].paragraph_format.alignment = WD_TAB_ALIGNMENT.DECIMAL
                    run = cell.paragraphs[0].add_run(str(data[rowid][cellid]))
                    run.font.name = '宋体'
                    run.font.size = 120000
        try:
            doc.save(word_path)
            return True
        except PermissionError:
            logger.exception('文件无法保存，请确认是否打开状态！')
            return False
        except Exception:
            logger.exception('未知错误！')
            return False


class ExcelFunctions:

    def __init__(self):
        pass

    # 读取excel文档到数据库，生成数据表(表字段col1...10)，数据库对象，默认Sheet1，从第n行读取
    def read_excel_to_db(self, file, table_name, db_object, row_start, sheet_name='Sheet1'):
        workbook = xlrd.open_workbook(file)
        sheet = workbook.sheet_by_name(sheet_name)
        # 先删表
        db_object.db_excute_sql('drop table if exists ' + table_name)
        # 生成建表语句
        sql = 'create table ' + table_name + ' ('
        for i in range(sheet.ncols): sql = sql + 'col' + str(i) + ' text,'
        sql = sql[:-1] + ')'
        db_object.db_excute_sql(sql)
        result = False
        for x in range(sheet.nrows - row_start):
            row = []
            for y in range(sheet.ncols):
                row.append(sheet.cell_value(x + row_start, y))
            # 数据入表
            result = db_object.db_excute_sql('insert into ' + table_name + ' values ' + str(tuple(row)))
        if result:
            return True
        else:
            return False

    # 把数据写入到excel文档
    def write_excel(self, file, data, row_start, sheet_name='Sheet1'):
        pass


class ControlFunctions:

    def __init__(self):
        self.pyautogui = pyautogui
        self.win32gui = win32gui
        self.win32con = win32con
        self.win32clipboard = win32clipboard

        # self.pyautogui.PAUSE = 0.1  # 为所有的PyAutoGUI函数增加延迟。默认延迟时间是0.1秒
        self.width, self.height = pyautogui.size()  # 获得屏幕的分辨率

    # 根据窗口名模糊搜索，返回句柄
    def find_window(self, name):
        all_windows = list()
        find_windows = list()

        def get_allwindow(hwnd, mouse):
            _windows = dict()
            if self.win32gui.IsWindow(hwnd) and self.win32gui.IsWindowEnabled(hwnd) and self.win32gui.IsWindowVisible(hwnd):
                if self.win32gui.GetWindowText(hwnd) and hwnd:
                    _windows[self.win32gui.GetWindowText(hwnd)] = hwnd
                    all_windows.append(_windows)

        self.win32gui.EnumWindows(get_allwindow, 0)
        for i in all_windows:
            for k, v in i.items():
                if name in k:
                    find_windows.append(i)
        if not len(find_windows):
            return False
        else:
            for i in find_windows[0].items(): return int(i[1])

    # 打开一个窗口，提供路径，参数默认为空
    def open_window(self, path, command=None):
        if not os.system('"' + path + '" ' + command):
            logger.info('打开 ' + '"' + path + '" ' + command + ' 成功！')
            return True

    # 打开默认浏览器
    def open_default_browser(self, url, new=0, autoraise=True):
        if webbrowser.open(url, new, autoraise):
            return True
        else:
            return False

    # 设置窗口前置、最大化
    def set_window(self, name, is_top, is_max):
        is_top = int(is_top)
        is_max = int(is_max)
        hwnd = self.find_window(name)
        if not hwnd:
            logger.exception('未找到窗口或找到多个重名窗口，请核查！')
            return False
        else:
            if is_top:
                self.win32gui.SetForegroundWindow(hwnd)  # 前置窗口
            if is_max:
                if not self.win32gui.ShowWindow(hwnd, self.win32con.SW_SHOWMAXIMIZED):  # 最大化
                    logger.exception('最大化失败！')
                    return False
                return True

    # 获取鼠标当前位置
    def get_mouse_position(self):
        x, y = self.pyautogui.position()
        return x, y

    # 设置鼠标移动位置，偏移x,y像素,duration是否瞬间移动到默认0
    def mouse_move_to(self, x, y, x_offset=0, y_offset=0, duration=0.25):
        self.pyautogui.moveTo(x + x_offset, y + y_offset)
        return True

    # 移动鼠标到某处并操作，包含单击、双击clicks=2，left鼠标左键、middle鼠标中键、right鼠标右键
    def mouse_move_click(self, x, y, x_offset=0, y_offset=0, button='left', _duration=0.25, clicks=1, interval=0.25):
        self.pyautogui.click(x + x_offset, y + y_offset, button=button, duration=_duration, clicks=1, interval=0.25)
        return True

    # 全屏找图，默认选择找到的第一张,返回找到多少张图片和第一张的坐标
    def screen_find(self, filename, chose=1):
        all_location = list()
        location = self.pyautogui.locateAllOnScreen(filename)
        for i in location: all_location.append(i)
        num = len(all_location)
        if not len(all_location):
            logger.exception('未找到图片位置！')
            return False
        else:
            position = all_location[chose - 1]
        return num, position

    # 屏幕截图，默认全屏，区域截屏region(左上角XY坐标值和宽度、高度)
    def screen_shot(self, filename=None, region=None):
        img = self.pyautogui.screenshot(filename, region=region)
        return img

    # 根据图片全屏查找，需提供宽度、高度、保存位置，坐标偏移
    def screen_shot_by_img(self, find_pic, width, height, savefile, x_offset=0, y_offset=0, chose=1):
        location = self.screen_find(find_pic, chose=chose)[1]
        if location:
            real_location = (location[0] + x_offset, location[1] + y_offset, width, height)
            self.screen_shot(filename=savefile, region=real_location)
            return True
        else:
            return False

    # 根据图片全屏查找，移动到图片中间位置
    def mouse_move_by_img(self, find_pic, chose=1):
        location = self.screen_find(find_pic, chose=chose)
        if location:
            x, y = self.pyautogui.center(location[1])
            self.mouse_move_to(x, y)
            return True
        else:
            return False

    # 输入字符串,不支持中文
    def write_string(self, word, interval=0.1):
        self.pyautogui.typewrite(word, interval)
        return True

    # 通过剪贴板输入中文
    def write_string_gbk(self, word):
        self.win32clipboard.OpenClipboard()
        self.win32clipboard.EmptyClipboard()
        self.win32clipboard.SetClipboardData(self.win32con.CF_UNICODETEXT, word)
        self.win32clipboard.CloseClipboard()
        return True

    '''
    ‘enter’(或‘return’ 或 ‘\n’) || 回车
    ‘esc’ || ESC键
    ‘shiftleft’, ‘shiftright’ || 左右SHIFT键
    ‘altleft’, ‘altright’ || 左右ALT键
    ‘ctrlleft’, ‘ctrlright’ || 左右CTRL键
    ‘tab’ (‘\t’) || TAB键
    ‘backspace’, ‘delete’ || BACKSPACE 、DELETE键
    ‘pageup’, ‘pagedown’ || PAGE UP 和 PAGE DOWN键
    ‘home’, ‘end’ || HOME 和 END键
    ‘up’, ‘down’, ‘left’,‘right’ || 箭头键
    ‘f1’, ‘f2’, ‘f3’…. || F1…….F12键
    ‘insert’ || INS或INSERT键
    ‘printscreen’ || PRTSC 或 PRINT SCREEN键
    ‘winleft’, ‘winright’ || Win键
    ‘command’ || Mac OS command键
    '''

    # 输入特殊按键
    def press_key(self, btn):
        self.pyautogui.press(btn)
        return True

    # 输入组合热键
    def press_hot_key(self, *key):
        if len(key) == 3:
            self.pyautogui.hotkey(key[0], key[1], key[2])
        elif len(key) == 2:
            self.pyautogui.hotkey(key[0], key[1])
        else:
            logger.exception('一个按键请使用press_key方法，最大支持3个组合键！')
            return False
        return True


class SSHFunctions:

    def __init__(self):
        self.shell = None
        self.data = None

    # 开始ssh操作，与ssh_exit配合
    def ssh_start(self, gate_ip, username, password, server_ip, port='22', verbose='0', log='0'):
        port = int(port)
        log = int(log)
        verbose = int(verbose)
        data = ""
        ssh = paramiko.SSHClient()
        ssh.load_system_host_keys()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy)
        try:
            ssh.connect(hostname=gate_ip, port=port, username=username, password=password)
        except Exception as e:
            logger.exception('连接ssh失败!')
            return False
        self.shell = ssh.invoke_shell()
        self.shell.settimeout(5)
        while True:
            try:
                x = self.shell.recv(1024)
            except Exception as e:
                return False
            time.sleep(0.5)
            if not x: break
            try:
                data += x.decode('utf-8')
            except Exception as e:
                data += x.decode('gbk')
            if verbose:
                logger.info(data)
                sys.stdout.write(data)
                sys.stdout.flush()
            if re.search('Select group:', data):
                self.shell.send("0\n")
                data = ""
            elif re.search('Select page:', data):
                self.shell.send("0\n")
                data = ""
            elif re.search('Select server:', data):
                text = [x.strip().split(':') for x in data.strip().replace('\r', '').split('\n') if ':' in x]
                text = [x[0] for x in text if server_ip in x[1]][0]
                self.shell.send("" + text + "\n")
                data = ""
            elif re.search('Select account:', data):
                text = [x.strip().split(':') for x in data.strip().replace('\r', '').split('\n') if ':' in x]
                text = [x[0] for x in text if 'root' in x[1]][0]
                self.shell.send("" + text + "\n")
                data = ""
            elif re.search('Warning', data):
                self.shell.send("\n")
                data = ""
            elif re.search('comment:', data):
                self.shell.send("\n")
                data = ""
            elif re.search('~]#', data):
                return True
        return False

    # 开始ssh操作，与ssh_exit配合(直连linux服务器,不需要选择服务器)
    def ssh_start_direct(self, server_ip, username, password, port='22', verbose='0', log='0'):
        port = int(port)
        log = int(log)
        verbose = int(verbose)
        data = ""
        ssh = paramiko.SSHClient()
        ssh.load_system_host_keys()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy)
        try:
            ssh.connect(hostname=server_ip, port=port, username=username, password=password)
        except Exception as e:
            logger.exception('连接ssh失败!')
            return False
        self.shell = ssh.invoke_shell()
        self.shell.settimeout(5)
        while True:
            try:
                x = self.shell.recv(1024)
            except Exception as e:
                return False
            time.sleep(0.5)
            if not x: break
            try:
                data += x.decode('utf-8')
            except Exception as e:
                data += x.decode('gbk')
            if verbose:
                logger.info(data)
                sys.stdout.write(data)
                sys.stdout.flush()
            if re.search('~]', data):
                return True
        return False

    # 获取ssh通道返回的信息，传入ssh对象、搜索字符串
    def ssh_search_recv(self, options, timeout=5, verbose='0'):
        verbose = int(verbose)
        timeout = int(timeout)
        data = ""
        self.shell.settimeout(timeout)
        while True:
            try:
                x = self.shell.recv(1024)
            except Exception as e:
                logger.exception('返回信息异常：' + e)
                return False
            time.sleep(0.1)
            if not x: break
            try:
                data += x.decode('utf-8')
            except Exception as e:
                data += x.decode('gbk')
            if verbose:
                logger.info(data)
                sys.stdout.write(data)
                sys.stdout.flush()
            if options:
                if re.search(options, data):
                    self.data = data.replace('[m', '').replace('[H', '').replace('[J', '').replace('[K', '').replace('[6;1H', '').replace('[7m',
                                                                                                                                                '').replace(
                        '[01;31m',
                        '').replace(
                        '[0m', '').replace('[m', '')
                    return True
            else:
                self.data = data.replace('[m', '').replace('[H', '').replace('[J', '').replace('[K', '').replace('[6;1H', '').replace('[7m', '').replace(
                    '[01;31m',
                    '').replace('[0m',
                                '').replace(
                    '[m', '')
                return True
        return False

    # 向ssh通道发送交互数据
    def ssh_send(self, param):
        self.data = ""
        self.shell.send(param + "\n")
        time.sleep(0.2)
        return True

    def ssh_ssrp_pack(self, param, options, savepath, timeout=5, verbose='0'):
        send = self.ssh_send(param)
        if send:
            recv = self.ssh_search_recv(options, timeout, verbose)
            if recv:
                if savepath:
                    pic = self.ssh_save_pic(savepath)
                    if pic:
                        return True
                    else:
                        return False
                else:
                    return True
            else:
                return False
        else:
            return False

    def ssh_exit(self):
        self.ssh_send('')
        if self.ssh_search_recv('~]'):
            self.ssh_send("exit")
            return True
        return False

    # 保存返回信息为图片
    def ssh_save_pic(self, savepath):
        try:
            font = ImageFont.truetype('static/lib/msyh.ttc', 18)
            size = [font.getsize(x) for x in self.data.strip().replace('\r', '').split('\n')]
            width = max(size)[0]
            # height = max([x[1] for x in size]) * (len(size) + 1)
            height = max([x[1] for x in size]) * len(size)

            img = Image.new('RGB', size=(width, height), color=(0, 0, 0))
            draw = ImageDraw.Draw(img)
            draw.text(xy=(0, 0), text=self.data.strip().replace('\r', ''), font=font)
            img.save(savepath)
            return True
        except Exception as e:
            logger.exception(e)
            return False

    # 分析硬盘空间使用情况df -hT
    def ssh_analysis_file_usage(self):
        data = [x for x in self.data.strip().replace('\r', '').split('\n')]
        data = [x for x in data[2:] if '%' in x]  # 只保留带%号的行
        data = [x.split(' ') for x in data]  # 根据空格进行拆分
        data = [[y for y in x if y != '' and y is not None][-2:] for x in data]  # 去除多余的空格，只取需要的字段
        for i in data: i[0] = int(i[0].replace('%', ''))  # 转换成int来比较
        data = max(data)  # 获得最大的
        if data[0] < 85:
            msg = '硬盘空间使用正常'
        elif data[0] > 85 and data[0] < 90:
            msg = '硬盘空间  ' + data[1] + ' 即将使用完毕，需处理！'
        else:
            msg = '警告，硬盘空间  ' + data[1] + ' 马上将满，需立即处理！'
        return msg

    def ssh_analysis_file_inode(self):
        data = [x for x in self.data.strip().replace('\r', '').split('\n')]
        data = [x for x in data[2:] if '%' in x]  # 只保留带%号的行
        data = [x.split(' ') for x in data]  # 根据空格进行拆分
        data = [[y for y in x if y != '' and y is not None][-2:] for x in data]  # 去除多余的空格，只取需要的字段
        for i in data: i[0] = int(i[0].replace('%', '').strip())  # 转换成int来比较
        data = max(data)  # 获得最大的
        if data[0] < 85:
            msg = '硬盘Inode使用正常'
        elif data[0] > 85 and data[0] < 90:
            msg = '硬盘Inode ' + data[1] + ' 即将使用完毕，需处理！'
        else:
            msg = '硬盘Inode ' + data[1] + ' 马上将满，需立即处理！'
        return msg

    def ssh_analysis_cpu_top(self):
        data = [x for x in self.data.strip().replace('\r', '').split('\n')]
        data = [x for x in data if 'Cpu' in x]
        data = [x.split(' ') for x in data]  # 根据空格进行拆分
        data = float([[y for y in x if y != '' and y is not None] for x in data][0][1].split('%')[0])  # 去除多余的空格，只取需要的字段
        if data < 85:
            msg = 'CPU使用率正常'
        elif data > 85 and data < 90:
            msg = 'CPU使用率异常，请处理！'
        else:
            msg = 'CPU使用率过高，需立即处理！'
        return msg

    def ssh_analysis_free(self):
        data = [x for x in self.data.strip().replace('\r', '').split('\n')]
        data = [x for x in data[2:] if 'Mem:' in x]  # 只保留带%号的行
        data = [x.split(' ') for x in data]  # 根据空格进行拆分
        data = [[y for y in x if y != '' and y is not None] for x in data][0]  # 去除多余的空格，只取需要的字段
        data = round(int(data[2].strip()) / int(data[1].strip()) * 100, 2)
        if data < 85:
            msg = '内存使用正常'
        elif data > 85 and data < 90:
            msg = '内存 ' + str(data) + '% 即将使用完毕，需处理！'
        else:
            msg = '警告，内存  ' + str(data) + '% 马上将满，需立即处理！'
        return msg

    def ssh_analysis_netstat(self, ports):
        ports = ports.split(',')
        msg = ''
        for i in ports:
            if i not in self.data:
                msg += i + '端口不通，请检查！'
        if not len(msg): msg = '网络端口正常'
        return msg

    def ssh_analysis_check_igate(self):
        if 'Restart' in self.data:
            msg = '服务总线故障，请核查！'
        else:
            msg = '服务总线运行正常'
        return msg

    def ssh_analysis_config_igate(self):
        day = datetime.datetime.now().weekday()  # 0-6 星期日-星期六
        if day == 4:
            result = self.ssh_ssrp_pack('sh /usr/local/service/script/backup/backup_config_igate.sh', '~]', None)
            if result:
                msg = '正常'
            else:
                msg = '异常'
        else:
            msg = '正常'
        return msg

    def ssh_analysis_db_size(self):
        data = [x for x in self.data.strip().replace('\r', '').split('\n')]
        print(data)


class LogicFunctions:
    # 逻辑操作
    pass


class DatabaseFunctions:

    def __init__(self, conn_str):
        self.conn_str = conn_str
        self.conn = sl.connect(self.conn_str)
        self.cursor = self.conn.cursor()

    # 执行sql查询语句
    def db_query_sql(self, sql):
        try:
            data = self.cursor.execute(sql)
            result = data.fetchall()
            return result
        except Exception as e:
            logger.exception('异常SQL：' + sql)

    # 执行update、delete、命令等
    def db_excute_sql(self, sql):
        cursor = self.conn.cursor()
        try:
            cursor.execute(sql)
        except Exception as e:
            pass
            logger.exception('异常SQL：' + sql)
        self.conn.commit()
        cursor.close()
        return True


class OtherFunctions:

    # 其他类操作
    def __init__(self):
        pass

    # 暂停 单位秒
    def time_stop(self, seconds):
        time.sleep(int(seconds))
        return True

    # 文件拷贝
    def file_copy(self, old_file, new_file):
        if os.path.isfile(old_file):
            try:
                if not os.path.exists(os.path.split(new_file)[0]):
                    os.makedirs(os.path.split(new_file)[0])
                shutil.copy(old_file, new_file)
                return True
            except Exception as e:
                logger.exception(e)
                return False
        else:
            return False


class WebServiceFunctions:
    def __init__(self):
        pass

    # 总线访问地址(不要?wsdl)
    def get_bus_status(self, url):
        xml_content = '''<?xml version="1.0" encoding="UTF-8"?>
        <SOAP-ENV:Envelope xmlns:SOAP-ENV="http://schemas.xmlsoap.org/soap/envelope/" xmlns:SOAP-ENC="http://schemas.xmlsoap.org/soap/encoding/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema">
        	<SOAP-ENV:Header>
        		<username>TBI-IGATE</username>
        		<password>TBI-IGATE</password>
        	</SOAP-ENV:Header>
        	<SOAP-ENV:Body>
        		<m:exec xmlns:m="http://soa.csg.cn" SOAP-ENV:encodingStyle="http://schemas.xmlsoap.org/soap/encoding/"/>
        	</SOAP-ENV:Body>
        </SOAP-ENV:Envelope>'''
        try:
            request = Request(url, data=xml_content.encode('utf-8'), headers={'Content-Type': 'text/xml'})
            res = urlopen(request)
            for i in res: xml = i

            root_node = etree.fromstring(text=xml)
            queue_list = []
            for body in root_node:
                for execResponse in body:
                    for queues in execResponse:
                        for queue in queues:
                            queue_dict = {}
                            for i in queue:
                                queue_dict[i.tag] = i.text
                            queue_list.append(queue_dict)
            return queue_list
        except Exception as e:
            logger.exception(e)
            return False

    def get_service_operation(self, url, begin_date, end_date):
        xml_content = '''<?xml version="1.0" encoding="UTF-8"?>
        <SOAP-ENV:Envelope xmlns:SOAP-ENV="http://schemas.xmlsoap.org/soap/envelope/" xmlns:SOAP-ENC="http://schemas.xmlsoap.org/soap/encoding/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xmlns:xsd="http://www.w3.org/2001/XMLSchema" xmlns:soa="http://soa.csg.cn">
          <SOAP-ENV:Body>
            <soa:exec>
              <soa:BEGIN_DATE soa:type="string">%s</soa:BEGIN_DATE>
              <soa:END_DATE soa:type="string">%s</soa:END_DATE>
            </soa:exec>
          </SOAP-ENV:Body>
        </SOAP-ENV:Envelope>
        ''' % (begin_date, end_date)
        try:
            request = Request(url, data=xml_content.encode('utf-8'), headers={'Content-Type': 'text/xml'})
            res = urlopen(request)
            for i in res: xml = i

            root_node = etree.fromstring(text=xml)
            ns = root_node.nsmap
            queue_list = []
            for body in root_node:
                for execResponse in body:
                    for response in execResponse:
                        for response2 in response:
                            queue_dict = {}
                            for i in response2:
                                for x in ns.values():
                                    i.tag = i.tag.replace('{' + x + '}', '')
                                queue_dict[i.tag] = i.text
                            queue_list.append(queue_dict)
            return queue_list
        except Exception as e:
            logger.exception(e)
            return False


class FwyxqkFunctions:
    def __init__(self):
        pass

    def read_fwyxqk_to_db(self, begin_date, end_date, data=None):
        fwyxqk.objects.filter(begin_date=begin_date, end_date=end_date).delete()
        for i in data:
            i['begin_date'] = begin_date
            i['end_date'] = end_date
            fwyxqk.objects.update_or_create(**i)
        return True

    def ywxtfw(self, begin_date, end_date):
        xtfw = ((1, '/1DEEB48F-A4F4-4E3D-8D02-6945B49BC163', '创建业扩工程电子化移交单'),
                (2, '/BA8D29FD-6AE6-49E5-9AE2-0A12E3F043D2', '发布客服报障工单'),
                (3, '/921187B5-2E6C-4440-BDBB-F322F0680642', '发布低压故障工单'),
                (4, '/A2BA0608-ACFB-457A-8A7D-21E9F981688B', '发布中压故障工单'),
                (5, '/F40439EA-BEBF-4C61-8B2B-7948CB38BD29', '同步工程项目概预算信息'),
                (6, '/F4EF625D-92CC-49E4-BE66-4F3132BB03BB', '同步工程合同信息和工程项目金任务分解（有合同部分）信息'),
                (7, '/595D6CB2-DE35-43DD-B4F2-A2A22B0841E8', '同步工程项目发票和资金任务分解（无合同部分）信息'),
                (8, '/FCD9C808-5703-499D-BF50-88191330D4E6', '同步物资·工程项目资金任务分解（无合同部分）信息'),
                (9, '/D7C200EA-7F45-4DF4-874E-B97649FB866F', '同步物资·工程项目资金任务分解（有合同部分）信息'),
                (10, '/DB721B9A-C2E3-498D-B59C-5F51F36F82D0', '同步物资发票信息'),
                (11, '/E0DD7175-0515-403B-836A-456B0EDC8680', '同步工程项目竣工验收设备资产清单信息'),
                (12, '/A21AA0DD-C7FC-4A53-9B16-5BCDFAD1ED0C', '同步工程项目结算报告（工程结算数据）信息'),
                (13, '/609C73B0-1E82-4166-AE54-4181742C7F5A', '同步报废设备清单'),
                (14, '/605EE22C-00A8-4B97-930C-2B66F8EC1FF6', '同步工资薪酬支付数据'),
                (15, '/C5677F19-B396-430D-A2E3-8F0055E36462', '同步住房公积金支付数据'),
                (16, '/DB2F4B8A-4D29-49AC-8605-3631B0847586', '同步专项定标结果/非招标采购结果'),
                (17, '/7B8F575F-CD27-4295-84DB-F5ECFB27159E', '同步停电检修申请单信息'),
                (18, '/BFDA3867-1849-4676-A9DA-0D07CB56CB5F', '同步停电检修申请执行信息'),
                (19, '/25084902-BB3E-4118-ADDB-17BEF5A306AD', '同步GIS功能位置'),
                (20, '/48B8C639-4DE8-4736-9B4F-5ED5E4C9B292', '同步设备台账信息变更单'),
                (21, '/BF2039E7-AAE5-44DB-A7D3-2A0768E36BF6', '同步企业年金支付数据'))
        fwyxqk_sort.objects.filter(begin_date=begin_date, end_date=end_date, type='2').delete()
        xtfw_dict = {}
        result = []
        for i in xtfw:
            try:
                fwmx = list(fwyxqk.objects.filter(service_code=i[1], begin_date=begin_date, end_date=end_date).values('zcs', 'sbs', 'cwxx'))[0]
                xtfw_dict['type'] = '2'
                xtfw_dict['begin_date'] = begin_date
                xtfw_dict['end_date'] = end_date
                xtfw_dict['no'] = i[0]
                xtfw_dict['service_code'] = i[1]
                xtfw_dict['service_name'] = i[2]
                xtfw_dict['dyl'] = fwmx['zcs']
                xtfw_dict['ycs'] = fwmx['sbs']
                if not fwmx['cwxx']:
                    xtfw_dict['zyyc'] = ''
                else:
                    cwxx = fwmx['cwxx'].split('||')
                    if len(cwxx) == 1:
                        xtfw_dict['zyyc'] = fwmx['cwxx']
                    else:
                        zyyc = []
                        for x, y in enumerate(cwxx):
                            err = re.findall(r'[^()]+', y)[-1]
                            try:
                                err_num = int(err)
                            except Exception:
                                err_num = 1
                            zyyc.append([err_num, x])
                        xtfw_dict['zyyc'] = cwxx[max(zyyc)[1]]
            except Exception:
                xtfw_dict['dyl'] = 0
                xtfw_dict['ycs'] = 0
                xtfw_dict['zyyc'] = ''
            fwyxqk_sort.objects.update_or_create(**xtfw_dict)
            xtfw_list = [xtfw_dict['no'], xtfw_dict['service_code'], xtfw_dict['service_name'], xtfw_dict['dyl'], xtfw_dict['ycs'], xtfw_dict['zyyc']]
            result.append(xtfw_list)  # 组合成返回数据用的列表
        return result

    def jzrz(self, begin_date, end_date):
        jzfw = ((1, '/671E30F0-741F-0130-D29A-005056B05DAA', '接收并分发集中认证数据主服务'),
                (2, '/48F8ACA0-7420-0130-D29B-005056B05DAA', '人员信息数据处理'))
        fwyxqk_sort.objects.filter(begin_date=begin_date, end_date=end_date, type='3').delete()
        jzfw_dict = {}
        result = []
        for i in jzfw:
            try:
                fwmx = list(fwyxqk.objects.filter(service_code=i[1], begin_date=begin_date, end_date=end_date).values('zcs', 'sbs', 'cwxx'))[0]
                jzfw_dict['type'] = '3'
                jzfw_dict['begin_date'] = begin_date
                jzfw_dict['end_date'] = end_date
                jzfw_dict['no'] = i[0]
                jzfw_dict['service_code'] = i[1]
                jzfw_dict['service_name'] = i[2]
                jzfw_dict['dyl'] = fwmx['zcs']
                jzfw_dict['ycs'] = fwmx['sbs']
                if not fwmx['cwxx']:
                    jzfw_dict['zyyc'] = ''
                else:
                    cwxx = fwmx['cwxx'].split('||')
                    if len(cwxx) == 1:
                        jzfw_dict['zyyc'] = fwmx['cwxx']
                    else:
                        zyyc = []
                        for x, y in enumerate(cwxx):
                            err = re.findall(r'[^()]+', y)[-1]
                            try:
                                err_num = int(err)
                            except Exception:
                                err_num = 1
                            zyyc.append([err_num, x])
                        jzfw_dict['zyyc'] = cwxx[max(zyyc)[1]]
            except Exception:
                jzfw_dict['dyl'] = 0
                jzfw_dict['ycs'] = 0
                jzfw_dict['zyyc'] = ''
            fwyxqk_sort.objects.update_or_create(**jzfw_dict)
            jzfw_list = [jzfw_dict['no'], jzfw_dict['service_code'], jzfw_dict['service_name'], jzfw_dict['dyl'], jzfw_dict['ycs'], jzfw_dict['zyyc']]
            result.append(jzfw_list)  # 组合成返回数据用的列表
        return result

    def exceptop20(self, begin_date, end_date):
        fwyxqk_sort.objects.filter(begin_date=begin_date, end_date=end_date, type='4').delete()
        etfw = list(
            fwyxqk.objects.filter(begin_date=begin_date, end_date=end_date).values('service_code', 'service_name', 'zcs', 'sbs', 'cwxx').order_by('-sbs')[:20])
        etfw_dict = {}
        result = []
        for x, y in enumerate(etfw):
            etfw_dict['type'] = '4'
            etfw_dict['begin_date'] = begin_date
            etfw_dict['end_date'] = end_date
            etfw_dict['no'] = x + 1
            etfw_dict['service_code'] = y['service_code']
            etfw_dict['service_name'] = y['service_name']
            etfw_dict['dyl'] = y['zcs']
            etfw_dict['ycs'] = y['sbs']
            if not y['cwxx']:
                etfw_dict['zyyc'] = ''
            else:
                cwxx = y['cwxx'].split('||')
                if len(cwxx) == 1:
                    etfw_dict['zyyc'] = y['cwxx']
                else:
                    zyyc = []
                    for x, y in enumerate(cwxx):
                        try:
                            err = re.findall(r'[^()]+', y)[-1]
                            err_num = int(err)
                        except Exception:
                            err_num = 1
                        zyyc.append([err_num, x])
                    etfw_dict['zyyc'] = cwxx[max(zyyc)[1]]
                fwyxqk_sort.objects.update_or_create(**etfw_dict)
                etfw_list = [etfw_dict['no'], etfw_dict['service_code'], etfw_dict['service_name'], etfw_dict['dyl'], etfw_dict['ycs'], etfw_dict['zyyc']]
                result.append(etfw_list)  # 组合成返回数据用的列表
        return result

    def gpfw(self, begin_date, end_date):
        fwyxqk_sort.objects.filter(begin_date=begin_date, end_date=end_date, type='5').delete()
        gpfw = list(
            fwyxqk.objects.filter(begin_date=begin_date, end_date=end_date).values('service_code', 'service_name', 'zcs', 'sbs', 'cwxx').order_by('-zcs')[:20])
        gpfw_dict = {}
        result = []
        for x, y in enumerate(gpfw):
            gpfw_dict['type'] = '5'
            gpfw_dict['begin_date'] = begin_date
            gpfw_dict['end_date'] = end_date
            gpfw_dict['no'] = x + 1
            gpfw_dict['service_code'] = y['service_code']
            gpfw_dict['service_name'] = y['service_name']
            gpfw_dict['dyl'] = y['zcs']
            gpfw_dict['ycs'] = y['sbs']
            if not y['cwxx']:
                gpfw_dict['zyyc'] = ''
            else:
                cwxx = y['cwxx'].split('||')
                if len(cwxx) == 1:
                    gpfw_dict['zyyc'] = y['cwxx']
                else:
                    zyyc = []
                    for x, y in enumerate(cwxx):
                        err = re.findall(r'[^()]+', y)[-1]
                        try:
                            err_num = int(err)
                        except Exception:
                            err_num = 1
                        zyyc.append([err_num, x])
                    gpfw_dict['zyyc'] = cwxx[max(zyyc)[1]]
            fwyxqk_sort.objects.update_or_create(**gpfw_dict)
            gpfw_list = [gpfw_dict['no'], gpfw_dict['service_code'], gpfw_dict['service_name'], gpfw_dict['dyl'], gpfw_dict['ycs'], gpfw_dict['zyyc']]
            result.append(gpfw_list)  # 组合成返回数据用的列表
        return result


# 内部函数！
# 自动换行功能，或根据分号换行
def autoWrap(text, segment=';'):
    if not text:
        return None
    else:
        result = ''
        text = text.replace('\n', '').replace('\r', '')
        if segment.isdigit():
            text_list = [text[x:x + int(segment)].strip() for x in range(0, len(text), int(segment))]
            for x in text_list[:-1]:
                result += x + '\n'
            result = result + text_list[-1]
        else:
            text_list = text.split(segment)
            if not len(text_list[-1]):
                text_list = text.split(segment)[:-1]
            if len(text_list) > 0:
                for y in text_list[:-1]:
                    result += y + segment + '\n'
                result = result + text_list[-1]
            else:
                result = text
        return result


def character_conversion(data):
    pass
